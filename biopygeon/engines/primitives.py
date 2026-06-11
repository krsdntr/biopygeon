import os
import requests
import json
from bs4 import BeautifulSoup
from Bio import SeqIO
from io import StringIO
import traceback

# Memory state for agent
AGENT_MEMORY = {}

def tool_store_variable(key: str, value: str) -> str:
    """Simpan variabel ke dalam memori jangka pendek agen."""
    AGENT_MEMORY[key] = value
    return f"Success: Variable '{key}' stored."

def tool_retrieve_variable(key: str) -> str:
    """Ambil variabel dari memori."""
    if key in AGENT_MEMORY:
        return str(AGENT_MEMORY[key])
    return f"Error: Variable '{key}' not found."

def tool_read_file(filepath: str) -> str:
    """Membaca isi file teks lokal."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"Error reading file: {e}"

def tool_write_file(filepath: str, content: str, append: bool = False) -> str:
    """Menyimpan atau menambahkan string ke dalam file lokal."""
    try:
        mode = "a" if append else "w"
        with open(filepath, mode, encoding="utf-8") as f:
            f.write(content)
        return f"Success: File saved to {filepath}"
    except Exception as e:
        return f"Error writing file: {e}"

def tool_http_request(url: str, method: str = "GET", payload: str = None) -> str:
    """Melakukan request HTTP API generic (misal ke EBI, NCBI eUtils, dll)."""
    try:
        if method.upper() == "GET":
            response = requests.get(url, timeout=30)
        else:
            json_payload = json.loads(payload) if payload else None
            response = requests.post(url, json=json_payload, timeout=30)
        response.raise_for_status()
        return response.text
    except Exception as e:
        return f"Error HTTP Request: {e}"

def tool_web_scrape(url: str) -> str:
    """Mengunduh konten teks mentah dari halaman web."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 Biopygeon/1.0'}
        res = requests.get(url, headers=headers, timeout=15)
        res.raise_for_status()
        soup = BeautifulSoup(res.text, "html.parser")
        return soup.get_text(separator="\n", strip=True)
    except Exception as e:
        return f"Error Scraping: {e}"

def tool_run_python(code: str) -> str:
    """
    Mengeksekusi skrip python dinamis. Dibatasi secara konseptual.
    Mengembalikan stdout (print output) atau error.
    """
    import sys
    from io import StringIO
    
    # Redirect stdout
    old_stdout = sys.stdout
    sys.stdout = mystdout = StringIO()
    
    # Execute code
    try:
        # Provide basic safe globals
        safe_globals = {
            "__builtins__": __builtins__,
            "print": print,
            "json": json,
        }
        exec(code, safe_globals)
        output = mystdout.getvalue()
        if not output.strip():
            output = "Code executed successfully without print output."
        return output
    except Exception as e:
        return f"Error executing Python code:\n{traceback.format_exc()}"
    finally:
        sys.stdout = old_stdout

def tool_fetch_sequence(gene_id: str, db: str = "uniprot") -> str:
    """Menarik sekuens FASTA dari ID spesifik."""
    try:
        if db.lower() == "uniprot":
            url = f"https://rest.uniprot.org/uniprotkb/{gene_id}.fasta"
            return tool_http_request(url)
        elif db.lower() == "ncbi":
            # NCBI Entrez EFETCH
            url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?db=protein&id={gene_id}&rettype=fasta&retmode=text"
            return tool_http_request(url)
        else:
            return f"Error: Database '{db}' not supported."
    except Exception as e:
        return f"Error fetching sequence: {e}"

def tool_convert_format(input_data: str, from_fmt: str, to_fmt: str) -> str:
    """Mengonversi format bioinformatika menggunakan Biopython SeqIO."""
    try:
        in_stream = StringIO(input_data)
        out_stream = StringIO()
        records = list(SeqIO.parse(in_stream, from_fmt.lower()))
        if not records:
            return "Error: No valid records found to convert."
        SeqIO.write(records, out_stream, to_fmt.lower())
        return out_stream.getvalue()
    except Exception as e:
        return f"Error converting format: {e}"

def tool_extract_text(filepath: str) -> str:
    """Mengekstrak teks dari PDF atau DOCX."""
    ext = filepath.split(".")[-1].lower()
    try:
        if ext == "pdf":
            import PyPDF2
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif ext == "docx":
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return "Error: Unsupported file type. Only .pdf and .docx are supported."
    except ImportError:
        return "Error: PyPDF2 or python-docx library is required."
    except Exception as e:
        return f"Error extracting text: {e}"

def tool_merge_documents(file_list_json: str, output_path: str) -> str:
    """Menggabungkan dokumen. file_list_json harus string representasi list path."""
    try:
        files = json.loads(file_list_json)
        if not isinstance(files, list):
            return "Error: file_list_json must be a JSON list of paths."
            
        exts = set([f.split(".")[-1].lower() for f in files])
        if "pdf" in exts and "docx" not in exts:
            import PyPDF2
            merger = PyPDF2.PdfMerger()
            for pdf in files:
                merger.append(pdf)
            merger.write(output_path)
            merger.close()
            return f"Success: PDF documents merged to {output_path}"
        elif "docx" in exts and "pdf" not in exts:
            from docx import Document
            merged_doc = Document()
            for f in files:
                doc = Document(f)
                for p in doc.paragraphs:
                    merged_doc.add_paragraph(p.text)
            merged_doc.save(output_path)
            return f"Success: DOCX documents merged to {output_path}"
        else:
            return "Error: Cannot mix format, or unsupported formats."
    except ImportError:
        return "Error: Library not found."
    except Exception as e:
        return f"Error merging documents: {e}"
